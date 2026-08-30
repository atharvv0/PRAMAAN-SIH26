from __future__ import annotations

from pathlib import Path
from uuid import uuid4

from docx import Document

from ..db.database import FileRecord, Project, Task, session_scope
from ..db.repository import UPLOAD_ROOT, repo


def generate_approval_note(task_id: str, title: str, intent: str, state: dict) -> dict:
    doc = Document()
    doc.add_heading("PRAMAAN Approval Note", level=0)
    doc.add_paragraph(title)
    doc.add_heading("Task Intent", level=1)
    doc.add_paragraph(intent)
    doc.add_heading("Execution Summary", level=1)
    completed = state.get("completed_steps", [])
    doc.add_paragraph(f"Completed steps: {len(completed)}")
    doc.add_heading("Evidence", level=1)
    evidence = state.get("evidence", [])
    if evidence:
        for item in evidence:
            p=doc.add_paragraph(style="List Bullet")
            p.add_run(item.get("claim", ""))
            src=item.get("source")
            if src:p.add_run(f" — source: {src}")
    else:
        doc.add_paragraph("No evidence records were captured in this run.")
    doc.add_heading("Final Output", level=1)
    final=state.get("final_output") or {}
    doc.add_paragraph(str(final)[:10000])
    doc.add_heading("Approval", level=1)
    doc.add_paragraph("Human approval is required before release of the final deliverable.")
    path = UPLOAD_ROOT / f"{uuid4()}_Approval_Note.docx"
    doc.save(path)
    data=path.read_bytes()
    # Reuse repository's project lookup and create a file row.
    with session_scope() as s:
        task_row = s.get(Task, task_id)
        user_id = task_row.created_by
        project_id = task_row.project_id
        rec=FileRecord(file_id=str(uuid4()),project_id=project_id,uploaded_by=user_id,filename=path.name,mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",size_bytes=len(data),storage_path=str(path),sha256=__import__('hashlib').sha256(data).hexdigest(),sensitivity_class=task_row.sensitivity_class)
        s.add(rec); s.commit(); s.refresh(rec)
        repo.create_deliverable(task_id,rec.file_id,"docx","pending")
        return {"file_id":rec.file_id,"name":rec.filename}


def _artifact_request(intent: str) -> tuple[str, str] | None:
    import re
    low = intent.lower()
    match = re.search(r"(?:give|generate|create|save|return|export)\D{0,30}\.(txt|md|json|csv|docx)", low)
    if match:
        ext = match.group(1)
    elif "deliverable" in low or "file" in low:
        ext = "txt"
    else:
        return None

    return ext, f"summary.{ext}"


def maybe_generate_task_artifact(task_id: str, state: dict, title: str, intent: str) -> dict | None:
    requested = _artifact_request(intent)
    if not requested:
        return None

    ext, filename = requested
    final = state.get("final_output") or {}
    response = final.get("response") if isinstance(final, dict) else None
    if not response:
        return None

    import hashlib
    from datetime import datetime, timezone
    from pathlib import Path

    if "10 line" in intent.lower() or "10-line" in intent.lower():
        lines = [line.strip() for line in str(response).splitlines() if line.strip()]
        response = "\n".join(lines[:10])
        while len(response.splitlines()) < 10:
            response += "\n"

    path = UPLOAD_ROOT / f"{uuid4()}_{filename}"
    if ext in {"txt", "md"}:
        path.write_text(str(response).strip() + "\n", encoding="utf-8")
        mime = "text/plain" if ext == "txt" else "text/markdown"
    elif ext == "json":
        import json
        path.write_text(json.dumps({"title": title, "intent": intent, "response": str(response)}, indent=2), encoding="utf-8")
        mime = "application/json"
    elif ext == "csv":
        import csv
        with path.open("w", newline="", encoding="utf-8") as fh:
            writer = csv.writer(fh)
            writer.writerow(["section", "content"])
            writer.writerow(["response", str(response)])
        mime = "text/csv"
    elif ext == "docx":
        doc = Document()
        doc.add_heading(title or "PRAMAAN Deliverable", level=0)
        doc.add_heading("Task", level=1)
        doc.add_paragraph(intent)
        doc.add_heading("Response", level=1)
        doc.add_paragraph(str(response))
        doc.save(path)
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        return None

    data = path.read_bytes()
    with session_scope() as s:
        task_row = s.get(Task, task_id)
        if task_row is None:
            return None
        rec = FileRecord(
            file_id=str(uuid4()),
            project_id=task_row.project_id,
            uploaded_by=task_row.created_by,
            filename=filename,
            mime_type=mime,
            size_bytes=len(data),
            storage_path=str(path),
            sha256=hashlib.sha256(data).hexdigest(),
            sensitivity_class=task_row.sensitivity_class,
        )
        s.add(rec)
        s.commit()
        repo.create_deliverable(task_id, rec.file_id, ext, "approved")
        return {"file_id": rec.file_id, "name": rec.filename, "type": ext}
